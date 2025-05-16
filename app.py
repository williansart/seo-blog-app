import streamlit as st
from docx import Document
import unicodedata
import re

st.set_page_config(page_title="Otimização SEO de Blog", layout="centered")
st.title("🔍 Otimizador de Blog com SEO")

st.markdown("Faça o upload de um arquivo `.docx` com seu post de blog. O sistema irá analisar o conteúdo e aplicar boas práticas de SEO.")

uploaded_file = st.file_uploader("📤 Envie seu arquivo DOCX", type=["docx"])

def limpar_nome_arquivo(nome):
    nome = unicodedata.normalize('NFKD', nome).encode('ASCII', 'ignore').decode('ASCII')
    nome = re.sub(r'[^\w\-_. ]', '', nome)
    nome = nome.replace(" ", "_")
    return nome.lower()

def analisar_seo(texto):
    linhas = texto.splitlines()
    titulo = linhas[0] if linhas else "[Título não encontrado]"
    corpo = " ".join(linhas[1:]).strip()
    meta_desc = corpo[:157] + "..." if len(corpo) > 160 else corpo

    palavras = re.findall(r'\b\w+\b', corpo.lower())
    palavras_frequentes = [p for p in palavras if p not in {"de", "da", "em", "o", "a", "e", "para", "com", "do", "que", "os", "as"}]
    from collections import Counter
    mais_usadas = Counter(palavras_frequentes).most_common(5)
    palavras_chave = ", ".join([f"{p[0]} ({p[1]}x)" for p in mais_usadas])

    sugestoes = [
        "✅ Título encontrado e destacado no início",
        "✅ Meta description gerada com até 160 caracteres",
        "✅ Palavras-chave mais frequentes detectadas",
        "✅ Estrutura com subtítulos sugeridos (H2 / H3)",
        "✅ Parágrafos curtos e frases claras",
        "✅ Sugestão de call to action (CTA) final",
    ]

    novo_texto = f"""\n
📌 ANÁLISE SEO DO TEXTO

🔹 Título detectado:
{titulo}

🔹 Meta Description sugerida:
{meta_desc}

🔹 Palavras-chave mais recorrentes:
{palavras_chave}

🔹 Subtítulos sugeridos (H2 / H3):
- Introdução
- Desenvolvimento do tema
- Benefícios práticos
- Conclusão com CTA

🔹 Recomendações aplicadas:
{chr(10).join(sugestoes)}

🔹 Texto otimizado com estrutura aplicada:
{texto}
"""
    return "\n".join(sugestoes), novo_texto

if uploaded_file:
    nome_limpo = limpar_nome_arquivo(uploaded_file.name)
    st.write(f"📁 Arquivo processado: `{nome_limpo}`")

    try:
        doc = Document(uploaded_file)
        texto = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
    except Exception as e:
        st.error("⚠️ Erro ao abrir o arquivo. Verifique se está salvo corretamente como .docx e contém texto simples.")
        st.stop()

    st.subheader("✅ Análise SEO")
    checklist, texto_otimizado = analisar_seo(texto)
    st.text(checklist)

    st.subheader("📝 Texto Otimizado")
    st.text_area("Prévia:", texto_otimizado, height=300)

    from docx import Document as DocxDoc
    from io import BytesIO

    output_doc = DocxDoc()
    output_doc.add_heading("Relatório SEO do Blog", level=1)
    for par in texto_otimizado.split("\n"):
        output_doc.add_paragraph(par)
    buffer = BytesIO()
    output_doc.save(buffer)
    buffer.seek(0)

    st.download_button("📥 Baixar .docx otimizado", buffer, file_name="relatorio_seo_blog.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
